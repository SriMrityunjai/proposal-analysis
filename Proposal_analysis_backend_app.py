# Proposal_analysis_backend_app.py
import io
import os
import json
import re
from typing import List, Optional, Tuple, Any
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
from dotenv import load_dotenv
import pandas as pd
import logging

# ---------- Load Environment & Logging ----------
load_dotenv()
logging.basicConfig(level=logging.INFO)

# ---------- Config ----------
OPENAI_API_KEY = os.getenv("OPENROUTER_API_KEY") or os.getenv("OPENAI_API_KEY")
OPENAI_BASE_URL = os.getenv("OPENROUTER_BASE_URL", os.getenv("OPENAI_BASE_URL", "https://openrouter.ai/api/v1"))
#OPENAI_MODEL = os.getenv("OPENAI_MODEL", os.getenv("OPENROUTER_MODEL", "tngtech/deepseek-r1t2-chimera:free"))
OPENAI_MODEL = os.getenv("OPENAI_MODEL", os.getenv("OPENROUTER_MODEL", "x-ai/grok-4.1-fast:free"))

if not OPENAI_API_KEY:
    logging.warning("OPENAI/OPENROUTER API key not found in environment variables (OPENROUTER_API_KEY or OPENAI_API_KEY).")

# ---------- Initialize FastAPI ----------
app = FastAPI(title="Proposal Analysis Backend API", version="4.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Prompt Loader ----------
def load_prompt(name: str) -> str:
    """
    Load a prompt file from prompts/<name>.txt.
    Raises FileNotFoundError if missing.
    """
    path = os.path.join("prompts", f"{name}.txt")
    if not os.path.exists(path):
        raise FileNotFoundError(f"Prompt file not found: {path}")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

# ---------- File text extraction helpers ----------
try:
    from docx import Document
except Exception:
    Document = None

try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

def extract_text_from_docx(file_bytes: bytes) -> str:
    if Document is None:
        raise ImportError('python-docx not installed')
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text_from_pdf(file_bytes: bytes) -> str:
    if PdfReader is None:
        raise ImportError('PyPDF2 not installed')
    reader = PdfReader(io.BytesIO(file_bytes))
    text = []
    for page in reader.pages:
        try:
            t = page.extract_text()
            if t:
                text.append(t)
        except Exception:
            continue
    return "\n".join(text)

def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode(errors='ignore')

def extract_text(filename: str, file_bytes: bytes) -> str:
    """Smartly select extraction by file extension."""
    fname = filename.lower()
    try:
        if fname.endswith('.pdf'):
            return extract_text_from_pdf(file_bytes)
        if fname.endswith('.docx'):
            return extract_text_from_docx(file_bytes)
        if fname.endswith('.txt'):
            return extract_text_from_txt(file_bytes)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from {filename}: {str(e)}")
    # fallback
    return file_bytes.decode(errors='ignore')

# ---------- OpenAI/OpenRouter client (with timeout) ----------
def make_openai_client(timeout: int = 300):
    if not OPENAI_API_KEY:
        raise RuntimeError("Missing OPENAI/OPENROUTER API key in environment (OPENROUTER_API_KEY or OPENAI_API_KEY).")
    return OpenAI(base_url=OPENAI_BASE_URL, api_key=OPENAI_API_KEY, timeout=timeout)

def call_chat_completion(prompt: str, system: Optional[str] = None, model: Optional[str] = None, max_tokens: int = 4000):
    """
    Calls the OpenAI/OpenRouter chat completion API and returns the assistant content string.
    """
    model_to_use = model or OPENAI_MODEL
    client = make_openai_client()
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    try:
        resp = client.chat.completions.create(model=model_to_use, messages=messages, max_tokens=max_tokens)
        return resp.choices[0].message.content
    except Exception as e:
        logging.exception("Chat completion call failed")
        raise

# ---------- JSON extraction & parsing helpers ----------
def extract_json_block(raw: str) -> Any:
    """
    Try to parse JSON from a code block ```json ... ``` or from raw text.
    Returns parsed JSON object.
    """
    try:
        # first try to find a fenced code block (json or not)
        m = re.search(r"```(?:json)?\s*(.*?)\s*```", raw, flags=re.S)
        candidate = m.group(1).strip() if m else raw.strip()
        return json.loads(candidate)
    except Exception as e:
        raise ValueError(f"Failed to parse JSON from model output: {str(e)}\nOutput preview: {raw[:1000]}")

def parse_two_json_arrays_back_to_back(raw: str) -> Tuple[List[Any], List[Any]]:
    """
    Some prompts (evaluation.txt) require returning two JSON arrays back-to-back without surrounding text.
    This function tries to extract two JSON arrays appearing one after another in raw output.
    Returns (first_array, second_array) or raises ValueError.
    """
    # attempt to find the first '[' then parse the first array, then parse the second
    s = raw.strip()
    # remove surrounding code fences if present
    if s.startswith("```"):
        # keep inner content
        m = re.search(r"```(?:json)?\s*(.*?)\s*```", s, flags=re.S)
        if m:
            s = m.group(1).strip()

    # find first JSON array
    first_array = None
    second_array = None
    # Use regex to find two top-level arrays
    pattern = r"(\[\s*(?:.|\s)*?\s*\])"
    matches = re.findall(pattern, s)
    if len(matches) >= 2:
        try:
            first_array = json.loads(matches[0])
            second_array = json.loads(matches[1])
            return first_array, second_array
        except Exception as e:
            raise ValueError(f"Failed parsing two JSON arrays: {str(e)}")
    # fallback: if entire text is single JSON (object containing two arrays) try parse
    try:
        parsed = json.loads(s)
        # check if parsed is list-of-two-lists or dict containing two lists
        if isinstance(parsed, list) and len(parsed) == 2 and isinstance(parsed[0], list) and isinstance(parsed[1], list):
            return parsed[0], parsed[1]
        # dict with keys "requirements" and "scoring"
        if isinstance(parsed, dict) and "requirements" in parsed and "scoring" in parsed:
            return parsed["requirements"], parsed["scoring"]
    except Exception:
        pass
    raise ValueError("Could not find two JSON arrays in the model output.")

# ---------- Prompt builders ----------
def build_summary_prompt(rfp_text: str) -> str:
    return load_prompt("summary").replace("{rfp_text}", rfp_text)

def build_rule_extraction_prompt(rfp_text: str) -> str:
    return load_prompt("rules").replace("{rfp_text}", rfp_text)

def build_scoring_extraction_prompt(rfp_text: str) -> str:
    return load_prompt("scoring").replace("{rfp_text}", rfp_text)

def build_evaluation_prompt(rules_json: str, scoring_json: str, proposal_text: str) -> str:
    tmpl = load_prompt("evaluation")
    return tmpl.replace("{rules_json}", rules_json).replace("{scoring_json}", scoring_json).replace("{proposal_text}", proposal_text)

def build_compare_prompt(csv_data: str, ranking_list: List[dict]) -> str:
    """
    Use the compare_evaluation prompt template file (compare_evaluation.txt).
    Template expects placeholders:
      {csv_data} and {ranking_list}
    """
    tmpl = load_prompt("compare_evaluation")
    return tmpl.replace("{csv_data}", csv_data).replace("{ranking_list}", json.dumps(ranking_list, indent=2))

# ---------- Small utility ----------
def safe_numeric(val) -> Optional[float]:
    """Attempt to extract numeric value from val which might be string like '8/10' or '8' or 8."""
    try:
        if val is None:
            return None
        if isinstance(val, (int, float)):
            return float(val)
        s = str(val).strip()
        # if pattern like '8/10' take numerator
        m = re.match(r"^\s*([0-9]+(?:\.[0-9]+)?)\s*/\s*([0-9]+(?:\.[0-9]+)?)\s*$", s)
        if m:
            return float(m.group(1))
        # if contains parentheses like "Yes (8/10)" or digits inside
        m2 = re.search(r"([0-9]+(?:\.[0-9]+)?)", s)
        if m2:
            return float(m2.group(1))
        return None
    except Exception:
        return None

# ---------- API Endpoints ----------

@app.post("/extract_rfp")
async def extract_rfp(file: UploadFile = File(...)):
    """
    Extract summary, rules and scoring JSON from an uploaded RFP document.
    Returns JSON object with keys: summary (plain text), rules (parsed JSON list), scoring (parsed JSON list), errors (optional)
    """
    errors = []
    try:
        file_bytes = await file.read()
        text = extract_text(file.filename, file_bytes)
    except Exception as e:
        logging.exception("Text extraction failed")
        return {"error": f"Text extraction failed: {str(e)}"}

    model = OPENAI_MODEL

    # summary
    try:
        summary_prompt = build_summary_prompt(text)
        summary = call_chat_completion(summary_prompt, system="Summarize RFP", model=model)
    except Exception as e:
        logging.exception("Summary extraction failed")
        summary = ""
        errors.append(f"Summary extraction failed: {str(e)}")

    # rules extraction
    try:
        rule_prompt = build_rule_extraction_prompt(text)
        rules_raw = call_chat_completion(rule_prompt, system="Extract rules JSON", model=model)
        rules = extract_json_block(rules_raw)
    except Exception as e:
        logging.exception("Rule extraction/parse failed")
        rules = []
        errors.append(f"Rule extraction/parse failed: {str(e)}")

    # scoring extraction
    try:
        scoring_prompt = build_scoring_extraction_prompt(text)
        scoring_raw = call_chat_completion(scoring_prompt, system="Extract scoring JSON", model=model)
        scoring = extract_json_block(scoring_raw)
    except Exception as e:
        logging.exception("Scoring extraction/parse failed")
        scoring = []
        errors.append(f"Scoring extraction/parse failed: {str(e)}")

    result = {"summary": summary.strip(), "rules": rules, "scoring": scoring}
    if errors:
        result["errors"] = errors
    return result

@app.post("/evaluate_proposal")
async def evaluate_proposal(
    file: UploadFile = File(...),
    rules_json: str = Form(None),
    scoring_json: str = Form(None)
):
    """
    Evaluate a single proposal using rules_json and scoring_json.
    Returns:
      - requirements: list of requirement evaluation dicts (Addressed Yes/No)
      - scoring: list of scoring evaluation dicts (Score_Awarded numeric expected)
      - requirements_csv, scoring_csv: CSV strings (full rows) ready for download
      - totals: total_points (numeric sum if available)
    """
    try:
        file_bytes = await file.read()
        proposal_text = extract_text(file.filename, file_bytes)
    except Exception as e:
        logging.exception("Proposal text extraction failed")
        return {"error": f"Failed to read or extract proposal text: {str(e)}"}

    if not rules_json or not scoring_json:
        return {"error": "Both rules_json and scoring_json are required"}

    # validate JSON strings
    try:
        _ = json.loads(rules_json)
        _ = json.loads(scoring_json)
    except Exception as e:
        return {"error": f"Invalid JSON for rules or scoring: {str(e)}"}

    # build prompt
    try:
        prompt = build_evaluation_prompt(rules_json, scoring_json, proposal_text[:20000])  # limit size
    except Exception as e:
        logging.exception("Failed to build evaluation prompt")
        return {"error": f"Failed to load evaluation prompt template: {str(e)}"}

    model = OPENAI_MODEL
    try:
        raw = call_chat_completion(prompt, system="Evaluate proposal; return two JSON arrays (requirements then scoring)", model=model, max_tokens=6000)
        # parse two JSON arrays back-to-back
        try:
            requirements_list, scoring_list = parse_two_json_arrays_back_to_back(raw)
        except Exception as e:
            # last resort: if model returned single array named 'requirements' or 'scoring' or object, try extract_json_block
            try:
                parsed = extract_json_block(raw)
                if isinstance(parsed, dict) and "requirements" in parsed and "scoring" in parsed:
                    requirements_list = parsed["requirements"]
                    scoring_list = parsed["scoring"]
                elif isinstance(parsed, list) and len(parsed) == 2 and isinstance(parsed[0], list) and isinstance(parsed[1], list):
                    requirements_list, scoring_list = parsed[0], parsed[1]
                else:
                    raise
            except Exception:
                logging.exception("Evaluation JSON parse failed")
                # fallback: return raw for inspection
                return {"error": "Failed to parse evaluation output into two JSON arrays", "raw_output": raw}
    except Exception as e:
        logging.exception("Evaluation call failed")
        return {"error": f"OpenAI/OpenRouter evaluation call failed: {str(e)}"}

    # Normalize requirements_list and scoring_list to ensure consistent fields
    # Requirements expected fields: RFP_Requirement_or_Criteria, Mandatory, Max_Points, Addressed, Explanation/Evidence
    req_rows = []
    for item in requirements_list:
        # Normalize keys (support a few possible key names)
        def get_field(d, *keys, default=""):
            for k in keys:
                if k in d:
                    return d[k]
            return default
        if not isinstance(item, dict):
            continue
        row = {
            "RFP_Requirement_or_Criteria": get_field(item, "RFP_Requirement_or_Criteria", "RFP_Requirement_or_Criteria", "RFP Requirement", "Requirement", "RFP_Requirement"),
            "Mandatory": get_field(item, "Mandatory", "IsMandatory", "Required", default=""),
            "Max_Points": get_field(item, "Max_Points", "Max Points", default=""),
            "Addressed": get_field(item, "Addressed", "Addressed?","Satisfied","Yes/No", default=""),
            "Explanation/Evidence": get_field(item, "Explanation/Evidence", "Explanation", "Evidence", default="")
        }
        req_rows.append(row)

    # Scoring expected fields: Point, Max_Points, Score_Awarded, Explanation/Evidence
    score_rows = []
    total_points = 0.0
    any_score = False
    for item in scoring_list:
        if not isinstance(item, dict):
            continue
        point = item.get("Point") or item.get("EvaluationCriteria") or item.get("Criterion") or item.get("Score Title") or ""
        max_points = item.get("Max_Points") or item.get("Max Points") or item.get("Score") or ""
        score_awarded = item.get("Score_Awarded") or item.get("Score") or item.get("Awarded") or ""
        explanation = item.get("Explanation/Evidence") or item.get("Explanation") or item.get("EvaluationDetail") or ""
        # attempt numeric extraction
        num = safe_numeric(score_awarded)
        if num is not None:
            total_points += num
            any_score = True
        score_rows.append({
            "Point": point,
            "Max_Points": max_points,
            "Score_Awarded": score_awarded,
            "Explanation/Evidence": explanation
        })

    # Convert to DataFrames and CSV strings for UI / download
    req_df = pd.DataFrame(req_rows)
    scoring_df = pd.DataFrame(score_rows)

    req_csv = req_df.to_csv(index=False)
    scoring_csv = scoring_df.to_csv(index=False)

    result = {
        "requirements": req_rows,
        "scoring": score_rows,
        "requirements_csv": req_csv,
        "scoring_csv": scoring_csv,
        "total_points": (total_points if any_score else None)
    }
    return result

@app.post("/compare_reports")
async def compare_reports(files: List[UploadFile] = File(...)):
    """
    Accept multiple evaluation JSON files (outputs from /evaluate_proposal).
    Each uploaded file can be either:
      - the new evaluate_proposal JSON (with 'requirements' and 'scoring' + CSV strings), OR
      - an older 'evaluations' JSON array (single list) -- we will try to interpret loosely.

    Returns:
      - comparison tables for requirements and scoring (as dicts and CSV strings)
      - ranking list (descending by total score or rules-count)
      - recommendation: LLM-produced explanation (JSON or raw text)
    """
    evals = []
    for f in files:
        try:
            raw = (await f.read()).decode(errors='ignore')
            content = json.loads(raw)
            evals.append({"filename": f.filename, "data": content})
        except Exception as e:
            logging.exception("Failed to read/parse evaluation file")
            return {"error": f"Failed to read/parse evaluation file {f.filename}: {str(e)}"}

    # We will build two comparison tables:
    # 1) rules_comparison: rows = requirement names, columns = filenames -> Addressed (Yes/No)
    # 2) scoring_comparison: rows = scoring Point names, columns = filenames -> Score_Awarded

    rules_combined = {}
    scoring_combined = {}

    # Collect totals and rule compliance counts
    totals = {}
    compliance_counts = {}

    for e in evals:
        fname = e['filename']
        data = e['data']
        totals[fname] = {"total": 0.0, "has_scores": False}
        compliance_counts[fname] = 0

        # Branch depending on shape of data
        if isinstance(data, dict) and "requirements" in data and "scoring" in data:
            requirements = data["requirements"]
            scoring = data["scoring"]
        elif isinstance(data, dict) and "evaluations" in data:
            # older format: try to split into requirements vs scoring based on presence of 'Addressed' and 'Score_Awarded'
            evaluations = data["evaluations"]
            # classify items that have 'Addressed' key as requirements else scoring if have Score/Points
            requirements = [it for it in evaluations if any(k.lower().startswith("address") for k in it.keys()) or "RFP_Requirement_or_Criteria" in it]
            scoring = [it for it in evaluations if any(k.lower().startswith("score") or k.lower().startswith("point") for k in it.keys())]
        else:
            # Try to infer if the file directly contains two arrays in top-level
            if isinstance(data, list) and len(data) == 2 and isinstance(data[0], list) and isinstance(data[1], list):
                requirements, scoring = data[0], data[1]
            else:
                # Unknown structure: attempt to find keys
                requirements = data.get("requirements", []) if isinstance(data, dict) else []
                scoring = data.get("scoring", []) if isinstance(data, dict) else []

        # Process requirements list
        for it in requirements:
            if not isinstance(it, dict):
                continue
            # normalize keys
            req_name = it.get("RFP_Requirement_or_Criteria") or it.get("RFP Requirement") or it.get("Requirement") or it.get("RFP_Requirement") or "Unknown Requirement"
            addressed = it.get("Addressed") if "Addressed" in it else it.get("Addressed?") if "Addressed?" in it else it.get("Satisfied") if "Satisfied" in it else ""
            addressed_str = str(addressed)
            # Count addressed yes
            if str(addressed_str).strip().lower() in ["yes", "true", "1"]:
                compliance_counts[fname] += 1

            if req_name not in rules_combined:
                rules_combined[req_name] = {}
            rules_combined[req_name][fname] = addressed_str

        # Process scoring list
        for it in scoring:
            if not isinstance(it, dict):
                continue
            point_name = it.get("Point") or it.get("EvaluationCriteria") or it.get("Criterion") or "Unknown Point"
            score_awarded = it.get("Score_Awarded") or it.get("Score") or it.get("Points") or ""
            if point_name not in scoring_combined:
                scoring_combined[point_name] = {}
            scoring_combined[point_name][fname] = score_awarded

            # numeric total accumulation
            num = safe_numeric(score_awarded)
            if num is not None:
                totals[fname]["total"] += num
                totals[fname]["has_scores"] = True

    # Build DataFrames
    rules_df = pd.DataFrame.from_dict(rules_combined, orient='index').fillna('')
    scoring_df = pd.DataFrame.from_dict(scoring_combined, orient='index').fillna('')

    rules_csv = rules_df.to_csv(index=True)
    scoring_csv = scoring_df.to_csv(index=True)

    # Ranking: prefer numeric totals if any file has numeric totals
    if any(v["has_scores"] for v in totals.values()):
        ranking_list = sorted(
            [{"filename": fn, "total_points": (vals["total"] or 0.0), "rule_compliance_count": compliance_counts.get(fn, 0)} for fn, vals in totals.items()],
            key=lambda x: x["total_points"], reverse=True
        )
    else:
        # fallback ranking by compliance count
        ranking_list = sorted(
            [{"filename": fn, "total_points": None, "rule_compliance_count": compliance_counts.get(fn, 0)} for fn in totals.keys()],
            key=lambda x: x["rule_compliance_count"], reverse=True
        )

    # Build compare prompt using compare_evaluation.txt prompt template
    try:
        compare_prompt = build_compare_prompt(rules_csv + "\n\n" + scoring_csv, ranking_list)
    except Exception as e:
        logging.exception("Failed to load compare prompt template")
        # fallback to a simple default prompt if file missing
        compare_prompt = (
            "You are an impartial RFP evaluator. Given the CSVs and ranking, provide concise unbiased explanation.\n\n"
            f"Rules CSV:\n{rules_csv}\n\nScoring CSV:\n{scoring_csv}\n\n"
            f"Computed Ranking (descending): {json.dumps(ranking_list, indent=2)}\n\n"
            "Return JSON like: {\"ranking\": [...], \"analysis\": \"...\"} OR plain text if JSON not possible."
        )

    try:
        model = OPENAI_MODEL
        rec_raw = call_chat_completion(compare_prompt, system="Provide unbiased evaluation explanation", model=model, max_tokens=2000)
        try:
            rec_json = extract_json_block(rec_raw)
        except Exception:
            rec_json = {"raw_output": rec_raw}
    except Exception as e:
        logging.exception("Comparison LLM call failed")
        rec_json = {"error": f"OpenAI/OpenRouter comparison call failed: {str(e)}"}

    result = {
        "rules_comparison_table": rules_df.to_dict(),
        "rules_csv": rules_csv,
        "scoring_comparison_table": scoring_df.to_dict(),
        "scoring_csv": scoring_csv,
        "ranking": ranking_list,
        "recommendation": rec_json
    }
    return result

# ---------- End of File ----------

# # Proposal_analysis_backend_app.py
# import io
# import os
# import json
# import re
# from typing import List, Optional, Tuple
# from fastapi import FastAPI, UploadFile, File, Form
# from fastapi.middleware.cors import CORSMiddleware
# from openai import OpenAI
# from dotenv import load_dotenv
# import pandas as pd
# import logging

# # ---------- Load Environment ----------
# load_dotenv()
# logging.basicConfig(level=logging.INFO)

# # ---------- Config ----------
# OPENAI_API_KEY = os.getenv("OPENROUTER_API_KEY") or os.getenv("OPENAI_API_KEY")
# OPENAI_BASE_URL = os.getenv("OPENROUTER_BASE_URL", os.getenv("OPENAI_BASE_URL", "https://openrouter.ai/api/v1"))
# OPENAI_MODEL = os.getenv("OPENAI_MODEL", os.getenv("OPENROUTER_MODEL", "tngtech/deepseek-r1t2-chimera:free"))

# if not OPENAI_API_KEY:
#     logging.warning("OPENAI/OPENROUTER API key not found in environment variables (OPENROUTER_API_KEY or OPENAI_API_KEY).")

# # ---------- Initialize ----------
# app = FastAPI(title="Proposal Analysis Backend API", version="4.0")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # ---------- Prompt Loader ----------
# def load_prompt(name: str) -> str:
#     path = os.path.join("prompts", f"{name}.txt")
#     if not os.path.exists(path):
#         raise FileNotFoundError(f"Prompt file not found: {path}")
#     with open(path, "r", encoding="utf-8") as f:
#         return f.read()

# # ---------- File text extraction ----------
# try:
#     from docx import Document
# except Exception:
#     Document = None

# try:
#     from PyPDF2 import PdfReader
# except Exception:
#     PdfReader = None

# def extract_text_from_docx(file_bytes: bytes) -> str:
#     if Document is None:
#         raise ImportError('python-docx not installed')
#     doc = Document(io.BytesIO(file_bytes))
#     return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     if PdfReader is None:
#         raise ImportError('PyPDF2 not installed')
#     reader = PdfReader(io.BytesIO(file_bytes))
#     text = []
#     for page in reader.pages:
#         try:
#             t = page.extract_text()
#             if t:
#                 text.append(t)
#         except Exception:
#             continue
#     return "\n".join(text)

# def extract_text_from_txt(file_bytes: bytes) -> str:
#     return file_bytes.decode(errors='ignore')

# def extract_text(filename: str, file_bytes: bytes) -> str:
#     fname = filename.lower()
#     try:
#         if fname.endswith('.pdf'):
#             return extract_text_from_pdf(file_bytes)
#         if fname.endswith('.docx'):
#             return extract_text_from_docx(file_bytes)
#         if fname.endswith('.txt'):
#             return extract_text_from_txt(file_bytes)
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract text from {filename}: {str(e)}")
#     return file_bytes.decode(errors='ignore')

# # ---------- OpenAI/OpenRouter client (with timeout) ----------
# def make_openai_client(timeout: int = 300):
#     if not OPENAI_API_KEY:
#         raise RuntimeError("Missing OPENAI/OPENROUTER API key in environment (OPENROUTER_API_KEY or OPENAI_API_KEY).")
#     return OpenAI(base_url=OPENAI_BASE_URL, api_key=OPENAI_API_KEY, timeout=timeout)

# def call_chat_completion(prompt: str, system: Optional[str] = None, model: Optional[str] = None, max_tokens: int = 4000):
#     model_to_use = model or OPENAI_MODEL
#     client = make_openai_client()
#     messages = []
#     if system:
#         messages.append({"role": "system", "content": system})
#     messages.append({"role": "user", "content": prompt})
#     try:
#         resp = client.chat.completions.create(model=model_to_use, messages=messages, max_tokens=max_tokens)
#         return resp.choices[0].message.content
#     except Exception as e:
#         logging.exception("Chat completion call failed")
#         raise

# def extract_json_block(raw: str):
#     """Try to parse JSON from a code block or raw text. Raise informative error on failure."""
#     try:
#         m = re.search(r"```(?:json)?\s*(.*?)\s*```", raw, flags=re.S)
#         candidate = m.group(1).strip() if m else raw.strip()
#         return json.loads(candidate)
#     except Exception as e:
#         raise ValueError(f"Failed to parse JSON from model output: {str(e)}\nOutput preview: {raw[:1000]}")

# # ---------- Prompt builders ----------
# def build_summary_prompt(rfp_text: str) -> str:
#     return load_prompt("summary").replace("{rfp_text}", rfp_text)

# def build_rule_extraction_prompt(rfp_text: str) -> str:
#     return load_prompt("rules").replace("{rfp_text}", rfp_text)

# def build_scoring_extraction_prompt(rfp_text: str) -> str:
#     return load_prompt("scoring").replace("{rfp_text}", rfp_text)

# def build_evaluation_prompt(rules_json: str, scoring_json: str, proposal_text: str) -> str:
#     tmpl = load_prompt("evaluation")
#     return tmpl.replace("{rules_json}", rules_json).replace("{scoring_json}", scoring_json).replace("{proposal_text}", proposal_text)

# # ---------- API Endpoints ----------
# @app.post("/extract_rfp")
# async def extract_rfp(file: UploadFile = File(...)):
#     """
#     Extract summary, rules and scoring JSON from an uploaded RFP document.
#     Returns JSON object with keys: summary, rules, scoring, errors (optional)
#     """
#     errors = []
#     try:
#         file_bytes = await file.read()
#         text = extract_text(file.filename, file_bytes)
#     except Exception as e:
#         logging.exception("Text extraction failed")
#         return {"error": f"Text extraction failed: {str(e)}"}

#     model = OPENAI_MODEL

#     # summary
#     try:
#         summary_prompt = build_summary_prompt(text)
#         summary = call_chat_completion(summary_prompt, system="Summarize RFP", model=model)
#     except Exception as e:
#         logging.exception("Summary extraction failed")
#         summary = ""
#         errors.append(f"Summary extraction failed: {str(e)}")

#     # rules
#     try:
#         rule_prompt = build_rule_extraction_prompt(text)
#         rules_raw = call_chat_completion(rule_prompt, system="Extract rules JSON", model=model)
#         try:
#             rules = extract_json_block(rules_raw)
#         except Exception as e:
#             logging.exception("Rule JSON parse failed")
#             rules = []
#             errors.append(f"Rule JSON parse failed: {str(e)}")
#     except Exception as e:
#         logging.exception("Rule extraction failed")
#         rules = []
#         errors.append(f"Rule extraction failed: {str(e)}")

#     # scoring
#     try:
#         scoring_prompt = build_scoring_extraction_prompt(text)
#         scoring_raw = call_chat_completion(scoring_prompt, system="Extract scoring JSON", model=model)
#         try:
#             scoring = extract_json_block(scoring_raw)
#         except Exception as e:
#             logging.exception("Scoring JSON parse failed")
#             scoring = []
#             errors.append(f"Scoring JSON parse failed: {str(e)}")
#     except Exception as e:
#         logging.exception("Scoring extraction failed")
#         scoring = []
#         errors.append(f"Scoring extraction failed: {str(e)}")

#     result = {"summary": summary.strip(), "rules": rules, "scoring": scoring}
#     if errors:
#         result["errors"] = errors
#     return result

# @app.post("/evaluate_proposal")
# async def evaluate_proposal(
#     file: UploadFile = File(...),
#     rules_json: str = Form(None),
#     scoring_json: str = Form(None)
# ):
#     """
#     Evaluate a single proposal using rules_json and scoring_json.
#     Expects evaluation prompt template at prompts/evaluation.txt with placeholders:
#       {rules_json}, {scoring_json}, {proposal_text}
#     Returns {"evaluations": [...] } where each item should include "Score" (numeric preferable).
#     """
#     try:
#         file_bytes = await file.read()
#         proposal_text = extract_text(file.filename, file_bytes)
#     except Exception as e:
#         logging.exception("Proposal text extraction failed")
#         return {"error": f"Failed to read or extract proposal text: {str(e)}"}

#     if not rules_json or not scoring_json:
#         return {"error": "Both rules_json and scoring_json are required"}

#     # validate JSON strings
#     try:
#         _ = json.loads(rules_json)
#         _ = json.loads(scoring_json)
#     except Exception as e:
#         return {"error": f"Invalid JSON for rules or scoring: {str(e)}"}

#     # build prompt
#     try:
#         prompt = build_evaluation_prompt(rules_json, scoring_json, proposal_text[:8000])
#     except Exception as e:
#         logging.exception("Failed to build evaluation prompt")
#         return {"error": f"Failed to load evaluation prompt template: {str(e)}"}

#     model = OPENAI_MODEL
#     try:
#         raw = call_chat_completion(prompt, system="Evaluate proposal; return strict JSON array", model=model, max_tokens=6000)
#         # try to parse JSON
#         try:
#             evaluations = extract_json_block(raw)
#             # ensure it's a list
#             if isinstance(evaluations, dict):
#                 evaluations = [evaluations]
#         except Exception:
#             # If parse fails, return raw output for inspection
#             logging.warning("Evaluation JSON parse failed; returning raw output")
#             evaluations = [{"raw_output": raw}]
#     except Exception as e:
#         logging.exception("Evaluation call failed")
#         return {"error": f"OpenAI/OpenRouter evaluation call failed: {str(e)}"}

#     return {"evaluations": evaluations}

# def safe_numeric(val) -> Optional[float]:
#     """Attempt to extract numeric value from val which might be string like '8/10' or '8' or 8."""
#     try:
#         if val is None:
#             return None
#         if isinstance(val, (int, float)):
#             return float(val)
#         s = str(val).strip()
#         # if pattern like '8/10' take numerator
#         m = re.match(r"^\s*([0-9]+(?:\.[0-9]+)?)\s*/\s*([0-9]+(?:\.[0-9]+)?)\s*$", s)
#         if m:
#             return float(m.group(1))
#         # if contains parentheses like "Yes (8/10)"
#         m2 = re.search(r"([0-9]+(?:\.[0-9]+)?)", s)
#         if m2:
#             return float(m2.group(1))
#         return None
#     except Exception:
#         return None

# @app.post("/compare_reports")
# async def compare_reports(files: List[UploadFile] = File(...)):
#     """
#     Accept multiple evaluation JSON files (outputs from /evaluate_proposal).
#     Returns:
#       - comparison_table: dict (requirement -> {filename: "Yes (score)"})
#       - csv: CSV string
#       - ranking: ordered list [{"filename", "total_points"}] (descending)
#       - recommendation: model-produced JSON or raw text explaining ranking & reasons
#     """
#     evals = []
#     for f in files:
#         try:
#             raw = (await f.read()).decode()
#             content = json.loads(raw)
#             evals.append({"filename": f.filename, "data": content})
#         except Exception as e:
#             logging.exception("Failed to read/parse evaluation file")
#             return {"error": f"Failed to read/parse evaluation file {f.filename}: {str(e)}"}

#     # Build comparison table
#     combined = {}
#     for e in evals:
#         for ev in e['data'].get('evaluations', []):
#             # accommodate common key names
#             req = ev.get('RFP Requirement/Criteria') or ev.get('RFP Requirement') or ev.get('Requirement') or "Unknown Requirement"
#             addr = ev.get('Addressed') or ev.get('Addressed?') or ev.get('Satisfied') or ''
#             score = ev.get('Score') or ev.get('Points') or ev.get('Max Points') or None
#             if req not in combined:
#                 combined[req] = {}
#             if score is not None:
#                 combined[req][e['filename']] = f"{addr} ({score})"
#             else:
#                 combined[req][e['filename']] = f"{addr}"

#     df = pd.DataFrame.from_dict(combined, orient='index').fillna('')
#     csv_data = df.to_csv(index=True)

#     # Compute total points per file (sum numeric Score occurrences)
#     totals = {}
#     for e in evals:
#         total = 0.0
#         any_score = False
#         for ev in e['data'].get('evaluations', []):
#             score_val = ev.get('Score') or ev.get('Points') or ev.get('Max Points') or None
#             num = safe_numeric(score_val)
#             if num is not None:
#                 total += num
#                 any_score = True
#         totals[e['filename']] = {"total": total if any_score else None, "has_scores": any_score}

#     # Build ranking (only include those with numeric totals; if none have numeric totals, fallback to count of Addressed Yes)
#     if any(v["has_scores"] for v in totals.values()):
#         ranking_list = sorted(
#             [{"filename": fn, "total_points": (vals["total"] or 0.0)} for fn, vals in totals.items()],
#             key=lambda x: x["total_points"], reverse=True
#         )
#     else:
#         # fallback: rank by count of Addressed == Yes
#         counts = {}
#         for e in evals:
#             cnt = sum(1 for ev in e['data'].get('evaluations', []) if str(ev.get('Addressed')).lower() in ['yes', 'true', '1'])
#             counts[e['filename']] = cnt
#         ranking_list = sorted([{"filename": k, "total_points": v} for k, v in counts.items()], key=lambda x: x["total_points"], reverse=True)

#     # Create prompt for LLM explaining ranking and reasons (concise unbiased)
#     prompt = (
#         "You are an impartial RFP evaluator. Given the following comparison CSV and the computed ranking,\n"
#         "provide a concise unbiased explanation (2-6 short paragraphs) that: 1) states the ranked list, 2) explains why the top proposal is best, "
#         "3) lists key strengths of the top proposal and weaknesses of the others, and 4) lists any uncertainties or missing score data.\n\n"
#         f"Comparison CSV:\n{csv_data}\n\n"
#         f"Computed Ranking (descending): {json.dumps(ranking_list, indent=2)}\n\n"
#         "Return JSON like: {\"ranking\": [...], \"analysis\": \"...\"} OR plain text if JSON not possible."
#     )

#     try:
#         model = OPENAI_MODEL
#         rec_raw = call_chat_completion(prompt, system="Provide unbiased evaluation explanation", model=model, max_tokens=2000)
#         try:
#             rec_json = extract_json_block(rec_raw)
#         except Exception:
#             rec_json = {"raw_output": rec_raw}
#     except Exception as e:
#         logging.exception("Comparison LLM call failed")
#         rec_json = {"error": f"OpenAI/OpenRouter comparison call failed: {str(e)}"}

#     result = {
#         "comparison_table": df.to_dict(),
#         "csv": csv_data,
#         "ranking": ranking_list,
#         "recommendation": rec_json
#     }
#     return result
# # # # ---------- End of File ----------




# # Proposal_analysis_backend_app.py
# import io
# import os
# import json
# import re
# from typing import List, Optional
# from fastapi import FastAPI, UploadFile, File, Form
# from fastapi.middleware.cors import CORSMiddleware
# from openai import OpenAI
# from dotenv import load_dotenv
# import pandas as pd
# import logging

# # ---------- Load Environment ----------
# load_dotenv()
# logging.basicConfig(level=logging.INFO)

# # ---------- Config ----------
# OPENAI_API_KEY = os.getenv("OPENROUTER_API_KEY") or os.getenv("OPENAI_API_KEY")
# OPENAI_BASE_URL = os.getenv("OPENROUTER_BASE_URL", os.getenv("OPENAI_BASE_URL", "https://openrouter.ai/api/v1"))
# OPENAI_MODEL = os.getenv("OPENAI_MODEL", os.getenv("OPENROUTER_MODEL", "tngtech/deepseek-r1t2-chimera:free"))  # default safe model

# if not OPENAI_API_KEY:
#     logging.warning("OPENAI/OPENROUTER API key not found in environment variables (OPENROUTER_API_KEY or OPENAI_API_KEY).")

# # ---------- Initialize ----------
# app = FastAPI(title="Proposal Analysis Backend API", version="3.0")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # ---------- Prompt Loader ----------
# def load_prompt(name: str) -> str:
#     path = os.path.join("prompts", f"{name}.txt")
#     if not os.path.exists(path):
#         raise FileNotFoundError(f"Prompt file not found: {path}")
#     with open(path, "r", encoding="utf-8") as f:
#         return f.read()

# # ---------- File text extraction ----------
# try:
#     from docx import Document
# except Exception:
#     Document = None

# try:
#     from PyPDF2 import PdfReader
# except Exception:
#     PdfReader = None

# def extract_text_from_docx(file_bytes: bytes) -> str:
#     if Document is None:
#         raise ImportError('python-docx not installed')
#     doc = Document(io.BytesIO(file_bytes))
#     return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     if PdfReader is None:
#         raise ImportError('PyPDF2 not installed')
#     reader = PdfReader(io.BytesIO(file_bytes))
#     text = []
#     for page in reader.pages:
#         try:
#             t = page.extract_text()
#             if t:
#                 text.append(t)
#         except Exception:
#             continue
#     return "\n".join(text)

# def extract_text_from_txt(file_bytes: bytes) -> str:
#     return file_bytes.decode(errors='ignore')

# def extract_text(filename: str, file_bytes: bytes) -> str:
#     fname = filename.lower()
#     try:
#         if fname.endswith('.pdf'):
#             return extract_text_from_pdf(file_bytes)
#         if fname.endswith('.docx'):
#             return extract_text_from_docx(file_bytes)
#         if fname.endswith('.txt'):
#             return extract_text_from_txt(file_bytes)
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract text from {filename}: {str(e)}")
#     return file_bytes.decode(errors='ignore')

# # ---------- OpenAI/OpenRouter client ----------
# # ---------- OpenAI/OpenRouter client with timeout ----------
# def make_openai_client():
#     if not OPENAI_API_KEY:
#         raise RuntimeError("Missing OPENAI/OPENROUTER API key in environment (OPENROUTER_API_KEY or OPENAI_API_KEY).")

#     # 300 seconds (5 min) timeout for long RFP documents
#     return OpenAI(
#         base_url=OPENAI_BASE_URL,
#         api_key=OPENAI_API_KEY,
#         timeout=300
#     )

# # ---------- OpenAI/OpenRouter client with timeout ----------
# def make_openai_client():
#     if not OPENAI_API_KEY:
#         raise RuntimeError("Missing OPENAI/OPENROUTER API key in environment (OPENROUTER_API_KEY or OPENAI_API_KEY).")

#     # 300 seconds (5 min) timeout for long RFP documents
#     return OpenAI(
#         base_url=OPENAI_BASE_URL,
#         api_key=OPENAI_API_KEY,
#         timeout=300
#     )


# def call_chat_completion(prompt: str, system: Optional[str] = None, model: Optional[str] = None, max_tokens: int = 4000):
#     model_to_use = model or OPENAI_MODEL
#     client = make_openai_client()

#     messages = []
#     if system:
#         messages.append({"role": "system", "content": system})
#     messages.append({"role": "user", "content": prompt})

#     try:
#         resp = client.chat.completions.create(
#             model=model_to_use,
#             messages=messages,
#             max_tokens=max_tokens
#         )
#         return resp.choices[0].message.content
#     except Exception as e:
#         logging.error(f"Chat completion error: {e}")
#         raise


# def extract_json_block(raw: str):
#     """Try to parse JSON from a code block or raw text. Raise informative error on failure."""
#     try:
#         m = re.search(r"```(?:json)?\s*(.*?)\s*```", raw, flags=re.S)
#         candidate = m.group(1).strip() if m else raw.strip()
#         return json.loads(candidate)
#     except Exception as e:
#         raise ValueError(f"Failed to parse JSON from model output: {str(e)}\nOutput preview: {raw[:800]}")

# # ---------- Prompt builders (use your prompt files) ----------
# def build_summary_prompt(rfp_text: str) -> str:
#     return load_prompt("summary").replace("{rfp_text}", rfp_text)

# def build_rule_extraction_prompt(rfp_text: str) -> str:
#     return load_prompt("rules").replace("{rfp_text}", rfp_text)

# def build_scoring_extraction_prompt(rfp_text: str) -> str:
#     return load_prompt("scoring").replace("{rfp_text}", rfp_text)

# def build_evaluation_prompt(rules: list, scoring: list, proposal_text: str) -> str:
#     """Load evaluation prompt template and inject rules/scoring/proposal."""
#     prompt_template = load_prompt("evaluation")  # ensure prompts/evaluation.txt exists
#     prompt = prompt_template.replace("{rules_json}", json.dumps(rules, indent=2))
#     prompt = prompt.replace("{scoring_json}", json.dumps(scoring, indent=2))
#     prompt = prompt.replace("{proposal_text}", proposal_text[:6000])  # cap to reasonable length
#     return prompt

# # ---------- API Endpoints ----------
# @app.post("/extract_rfp")
# async def extract_rfp(file: UploadFile = File(...)):
#     """
#     Extract summary, rules and scoring JSON from an uploaded RFP document.
#     Returns JSON object with keys: summary (string), rules (list/dict), scoring (list/dict), errors (optional)
#     """
#     errors = []
#     try:
#         file_bytes = await file.read()
#         text = extract_text(file.filename, file_bytes)
#     except Exception as e:
#         return {"error": f"Text extraction failed: {str(e)}"}

#     model = OPENAI_MODEL
#     # summary
#     try:
#         summary_prompt = build_summary_prompt(text)
#         summary = call_chat_completion(summary_prompt, system="You summarize RFP documents.", model=model)
#     except Exception as e:
#         summary = ""
#         errors.append(f"Summary extraction failed: {str(e)}")

#     # rules
#     try:
#         rule_prompt = build_rule_extraction_prompt(text)
#         rules_raw = call_chat_completion(rule_prompt, system="Extract RFP rules and return JSON.", model=model)
#         try:
#             rules = extract_json_block(rules_raw)
#         except Exception as e:
#             logging.error("Rule JSON parse failed: %s", e)
#             rules = []
#             errors.append(f"Rule JSON parse failed: {str(e)}")
#     except Exception as e:
#         rules = []
#         errors.append(f"Rule extraction failed: {str(e)}")

#     # scoring
#     try:
#         scoring_prompt = build_scoring_extraction_prompt(text)
#         scoring_raw = call_chat_completion(scoring_prompt, system="Extract scoring details and return JSON.", model=model)
#         try:
#             scoring = extract_json_block(scoring_raw)
#         except Exception as e:
#             logging.error("Scoring JSON parse failed: %s", e)
#             scoring = []
#             errors.append(f"Scoring JSON parse failed: {str(e)}")
#     except Exception as e:
#         scoring = []
#         errors.append(f"Scoring extraction failed: {str(e)}")

#     result = {"summary": summary.strip(), "rules": rules, "scoring": scoring}
#     if errors:
#         result["errors"] = errors
#     return result

# @app.post("/evaluate_proposal")
# async def evaluate_proposal(
#     file: UploadFile = File(...),
#     rules_json: str = Form(None),
#     scoring_json: str = Form(None)
# ):
#     """
#     Evaluate a single proposal using rules_json and scoring_json (both required).
#     Returns {"evaluations": [...] } where evaluations is a list of per-requirement objects.
#     """
#     try:
#         file_bytes = await file.read()
#         proposal_text = extract_text(file.filename, file_bytes)
#     except Exception as e:
#         return {"error": f"Failed to read or extract proposal text: {str(e)}"}

#     if not rules_json or not scoring_json:
#         return {"error": "Both rules_json and scoring_json are required"}

#     try:
#         rules = json.loads(rules_json)
#         scoring = json.loads(scoring_json)
#     except Exception as e:
#         return {"error": f"Invalid JSON for rules or scoring: {str(e)}"}

#     # Build evaluation prompt from prompts/evaluation.txt and inject inputs
#     try:
#         evaluation_prompt_template = load_prompt("evaluation")
#     except Exception as e:
#         return {"error": f"Failed to load evaluation prompt template: {str(e)}"}

#     # Build the final prompt: the evaluation prompt should expect placeholders:
#     prompt = evaluation_prompt_template.replace("{rules_json}", json.dumps(rules, indent=2))
#     prompt = prompt.replace("{scoring_json}", json.dumps(scoring, indent=2))
#     prompt = prompt.replace("{proposal_text}", proposal_text[:6000])

#     model = OPENAI_MODEL
#     try:
#         client = make_openai_client()
#         response = client.chat.completions.create(
#             model=model,
#             messages=[
#                 {"role": "system", "content": "You are an expert RFP evaluator. Be objective and unbiased."},
#                 {"role": "user", "content": prompt}
#             ],
#             max_tokens=4000
#         )
#         raw_output = response.choices[0].message.content.strip()
#         try:
#             evaluations = json.loads(raw_output)
#         except Exception:
#             # If JSON parse fails, return raw output under 'raw_output' so frontend can show it.
#             evaluations = [{"raw_output": raw_output}]
#     except Exception as e:
#         return {"error": f"OpenAI/OpenRouter call failed: {str(e)}"}

#     return {"evaluations": evaluations}

# @app.post("/compare_reports")
# async def compare_reports(files: List[UploadFile] = File(...)):
#     """
#     Accept multiple evaluation JSON files (the outputs from /evaluate_proposal).
#     Returns:
#       - comparison_table: dict (aggregate table)
#       - csv: CSV string
#       - recommendation: LLM reasoning & final ranked list
#     """
#     evals = []
#     for f in files:
#         try:
#             raw = (await f.read()).decode()
#             content = json.loads(raw)
#             evals.append({"filename": f.filename, "data": content})
#         except Exception as e:
#             return {"error": f"Failed to read/parse evaluation file {f.filename}: {str(e)}"}

#     # Build comparison table (requirement x filename -> Addressed/score)
#     combined = {}
#     for e in evals:
#         for ev in e['data'].get('evaluations', []):
#             # accommodate different key naming
#             req = ev.get('RFP Requirement/Criteria') or ev.get('RFP Requirement') or ev.get('Requirement') or "Unknown Requirement"
#             addr = ev.get('Addressed') or ev.get('Addressed?', '') or ev.get('Addressed (Yes/No)', '')
#             score = ev.get('Score') or ev.get('Points') or None
#             # store both addressed and score if present
#             if req not in combined:
#                 combined[req] = {}
#             # Prefer a combined string like "Yes (8/10)" or "No"
#             if score is not None:
#                 combined[req][e['filename']] = f"{addr} ({score})"
#             else:
#                 combined[req][e['filename']] = f"{addr}"

#     df = pd.DataFrame.from_dict(combined, orient='index')
#     csv_data = df.to_csv(index=True)

#     # Build a human-friendly summary and call LLM to reason which proposal is best
#     # Create a concise prompt with the table and ask for unbiased ranking and reasons
#     table_preview = df.fillna("").to_dict()  # small representation
#     prompt = (
#         "You are an impartial evaluator. Given the following evaluation table (requirement -> proposal -> Addressed/Score),\n"
#         "produce a ranked list of proposals from best to worst, provide total points for each proposal (if score info present, sum them), "
#         "and give concise unbiased reasons why the top proposal is better than others. Also list weaknesses of the top proposal.\n\n"
#         f"Table (as CSV):\n{csv_data}\n\n"
#         "Output JSON with keys: rankings (ordered list of {name, total_points (if available), reason}), "
#         "analysis (text). Be concise and unbiased."
#     )

#     try:
#         model = OPENAI_MODEL
#         client = make_openai_client()
#         response = client.chat.completions.create(
#             model=model,
#             messages=[
#                 {"role": "system", "content": "You are an objective RFP comparison expert."},
#                 {"role": "user", "content": prompt}
#             ],
#             max_tokens=2000
#         )
#         comp_raw = response.choices[0].message.content.strip()
#         try:
#             comp_json = json.loads(comp_raw)
#         except Exception:
#             # Return raw LLM response if JSON parse fails
#             comp_json = {"raw_output": comp_raw}
#     except Exception as e:
#         comp_json = {"error": f"OpenAI/OpenRouter comparison call failed: {str(e)}"}

#     return {"comparison_table": df.to_dict(), "csv": csv_data, "recommendation": comp_json}
# # """
# # FastAPI backend for RFP Proposal Analysis POC
# # --------------------------------------------
# # Endpoints:
# # - POST /extract_rfp: Extracts summary, rules, and scoring sections from RFP.
# # - POST /evaluate_proposal: Evaluates proposal documents against RFP rules and scoring.
# # - POST /compare_reports: Compares multiple evaluation reports.

# # Run backend:
# #     pip install fastapi uvicorn python-docx PyPDF2 openai python-dotenv pandas
# #     uvicorn Proposal_analysis_backend_app:app --reload --port 8000

# # This backend interacts with the Streamlit frontend.
# # Ensure your OpenRouter API key is set:
# #     export OPENROUTER_API_KEY="sk-or-v1-..."
# # """

# # import io
# # import os
# # import json
# # import re
# # from typing import List
# # from fastapi import FastAPI, UploadFile, File, Form
# # from fastapi.middleware.cors import CORSMiddleware
# # from openai import OpenAI
# # from dotenv import load_dotenv
# # import pandas as pd

# # # ---------- Load Environment ----------
# # load_dotenv()

# # # ---------- Initialize ----------
# # app = FastAPI(title="Proposal Analysis Backend API", version="2.0")

# # app.add_middleware(
# #     CORSMiddleware,
# #     allow_origins=["*"],
# #     allow_credentials=True,
# #     allow_methods=["*"],
# #     allow_headers=["*"],
# # )

# # # ---------- Prompt Loader ----------
# # def load_prompt(name: str) -> str:
# #     """Load prompt template from prompts/ directory"""
# #     path = os.path.join("prompts", f"{name}.txt")
# #     with open(path, "r", encoding="utf-8") as f:
# #         return f.read()

# # # ---------- Helper Functions ----------
# # try:
# #     from docx import Document
# # except Exception:
# #     Document = None

# # try:
# #     from PyPDF2 import PdfReader
# # except Exception:
# #     PdfReader = None

# # def extract_text_from_docx(file_bytes: bytes) -> str:
# #     if Document is None:
# #         raise ImportError('python-docx not installed')
# #     doc = Document(io.BytesIO(file_bytes))
# #     return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# # def extract_text_from_pdf(file_bytes: bytes) -> str:
# #     if PdfReader is None:
# #         raise ImportError('PyPDF2 not installed')
# #     reader = PdfReader(io.BytesIO(file_bytes))
# #     text = []
# #     for page in reader.pages:
# #         try:
# #             t = page.extract_text()
# #             if t:
# #                 text.append(t)
# #         except Exception:
# #             continue
# #     return "\n".join(text)

# # def extract_text_from_txt(file_bytes: bytes) -> str:
# #     return file_bytes.decode(errors='ignore')

# # def extract_text(filename: str, file_bytes: bytes) -> str:
# #     fname = filename.lower()
# #     if fname.endswith('.pdf'):
# #         return extract_text_from_pdf(file_bytes)
# #     if fname.endswith('.docx'):
# #         return extract_text_from_docx(file_bytes)
# #     if fname.endswith('.txt'):
# #         return extract_text_from_txt(file_bytes)
# #     return file_bytes.decode(errors='ignore')

# # # ---------- OpenRouter Client ----------
# # def make_openrouter_client():
# #     base_url = os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
# #     api_key = os.getenv("OPENROUTER_API_KEY")
# #     if not api_key:
# #         raise RuntimeError("Missing OPENROUTER_API_KEY in environment")
# #     return OpenAI(base_url=base_url, api_key=api_key)

# # def call_openrouter_chat(prompt: str, model: str, system: str = None):
# #     client = make_openrouter_client()
# #     messages = []
# #     if system:
# #         messages.append({"role": "system", "content": system})
# #     messages.append({"role": "user", "content": prompt})
# #     completion = client.chat.completions.create(model=model, messages=messages)
# #     try:
# #         return completion.choices[0].message.content
# #     except Exception:
# #         return str(completion)

# # # ---------- Prompt Builders ----------
# # def build_summary_prompt(rfp_text: str) -> str:
# #     return load_prompt("summary").replace("{rfp_text}", rfp_text)

# # def build_rule_extraction_prompt(rfp_text: str) -> str:
# #     return load_prompt("rules").replace("{rfp_text}", rfp_text)

# # def build_scoring_extraction_prompt(rfp_text: str) -> str:
# #     return load_prompt("scoring").replace("{rfp_text}", rfp_text)

# # # ---------- API Endpoints ----------
# # @app.post("/extract_rfp")
# # async def extract_rfp(file: UploadFile = File(...)):
# #     file_bytes = await file.read()
# #     text = extract_text(file.filename, file_bytes)

# #     model = "deepseek/deepseek-r1-distill-llama-70b:free"
# #     summary_prompt = build_summary_prompt(text)
# #     rule_prompt = build_rule_extraction_prompt(text)
# #     scoring_prompt = build_scoring_extraction_prompt(text)

# #     summary = call_openrouter_chat(summary_prompt, model)
# #     rules_raw = call_openrouter_chat(rule_prompt, model)
# #     scoring_raw = call_openrouter_chat(scoring_prompt, model)

# #     def extract_json_block(raw):
# #         m = re.search(r"```(?:json)?\s*(.*?)\s*```", raw, flags=re.S)
# #         return json.loads(m.group(1)) if m else json.loads(raw)

# #     try:
# #         rules = extract_json_block(rules_raw)
# #     except Exception:
# #         rules = []

# #     try:
# #         scoring = extract_json_block(scoring_raw)
# #     except Exception:
# #         scoring = []

# #     return {"summary": summary.strip(), "rules": rules, "scoring": scoring}

# # @app.post("/evaluate_proposal")
# # async def evaluate_proposal(
# #     file: UploadFile = File(...),
# #     rules_json: str = Form(None),
# #     scoring_json: str = Form(None)
# # ):
# #     try:
# #         file_bytes = await file.read()
# #         proposal_text = extract_text(file.filename, file_bytes)
# #     except Exception as e:
# #         return {"error": f"Failed to read or extract text: {str(e)}"}

# #     if not rules_json or not scoring_json:
# #         return {"error": "Both rules_json and scoring_json are required"}

# #     try:
# #         rules = json.loads(rules_json)
# #         scoring = json.loads(scoring_json)
# #     except Exception as e:
# #         return {"error": f"Invalid JSON input: {str(e)}"}

# #     model = "deepseek/deepseek-r1-distill-llama-70b:free"
# #     client = make_openrouter_client()

# #     prompt = f"""
# # You are an expert government RFP evaluator.

# # You will be given:
# # 1. The RFP requirements and scoring criteria (in JSON)
# # 2. A proposal text submitted by a firm.

# # Evaluate each requirement and score item.
# # For each one, check if it is addressed in the proposal, extract supporting evidence, and
# # produce output in JSON format as follows:

# # [
# #   {{
# #     "RFP Requirement/Criteria": "...",
# #     "Max Points": "...",
# #     "Addressed": "Yes/No",
# #     "Explanation/Evidence from Proposal": "..."
# #   }}
# # ]

# # Be precise, objective, and quote or summarize directly from the proposal text.
# # If evidence is missing, mark "Addressed" as "No" and leave explanation empty.

# # RFP Rules and Criteria:
# # {json.dumps(rules, indent=2)}

# # RFP Scoring:
# # {json.dumps(scoring, indent=2)}

# # Proposal Text:
# # {proposal_text[:4000]}
# # """

# #     response = client.chat.completions.create(
# #         model=model,
# #         messages=[
# #             {"role": "system", "content": "You are an expert RFP evaluator."},
# #             {"role": "user", "content": prompt}
# #         ],
# #     )

# #     try:
# #         eval_text = response.choices[0].message.content.strip()
# #         evaluations = json.loads(eval_text)
# #     except Exception:
# #         evaluations = [{"raw_output": response.choices[0].message.content.strip()}]

# #     return {"evaluations": evaluations}

# # @app.post("/compare_reports")
# # async def compare_reports(files: List[UploadFile] = File(...)):
# #     evals = []
# #     for f in files:
# #         content = json.loads((await f.read()).decode())
# #         evals.append({"filename": f.filename, "data": content})

# #     combined = {}
# #     for e in evals:
# #         for ev in e['data'].get('evaluations', []):
# #             req = ev.get('RFP Requirement/Criteria', 'Unknown')
# #             sat = ev.get('Addressed', 'Unknown')
# #             if req not in combined:
# #                 combined[req] = {}
# #             combined[req][e['filename']] = sat

# #     df = pd.DataFrame.from_dict(combined, orient='index')
# #     csv_data = df.to_csv(index=True)

# #     return {"comparison_table": df.to_dict(), "csv": csv_data}
# # # ---------- End of File ----------